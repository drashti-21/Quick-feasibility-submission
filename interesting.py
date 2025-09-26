import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

def create_word_table_from_dict(data_dict):
    doc = Document()
    doc.add_heading("Data Table", level=1)
    table = doc.add_table(rows=0, cols=2, style="Table Grid")
   
    for key, value in data_dict.items():
        row = table.add_row()
        row.cells[0].text = str(key)
        row.cells[1].text = str(value)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def main():
    st.title("Quick Feasibility Submission:")
    st.write("Enter the details below:")

    # Input fields
    platform = st.text_input("Platform")
    domain = st.text_input("Domain")

    # Dropdown for Method
    method = st.selectbox(
        "Method",
        options=["REQUEST", "BROWSER"]
    )

    # Dropdown for Complexity
    complexity = st.selectbox(
        "Complexity",
        options=["LOW", "MEDIUM", "HIGH"]
    )

    # Dropdown for Proxy
    proxy = st.selectbox(
        "Proxy",
        options=["None","SCRAPDO", "SCRAPERAPI","DECODO"]
    )

    credit = st.text_input("Credit")
    request = st.text_input("Request")
    # Generate document on button click
    if st.button("Generate Document"):
        if not all([platform, domain, method, complexity, proxy, credit]):
            st.error("Please fill in all fields.")
        else:
            data = {
                "Platform": platform,
                "Domain": domain,
                "Method": method,
                "Complexity": complexity,
                "Proxy": proxy,
                "Credit": credit,
                "Request": request
            }
            output_filename = f"{platform}_doc.docx"
            bio = create_word_table_from_dict(data)
            st.download_button(
                label="Download Document",
                data=bio,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()



